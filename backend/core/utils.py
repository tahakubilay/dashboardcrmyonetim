import os
from datetime import datetime
from decimal import Decimal
from django.conf import settings
from django.core.files.base import ContentFile
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from io import BytesIO
import pandas as pd


def generate_unique_filename(prefix, extension):
    """Benzersiz dosya adı oluştur"""
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"{prefix}_{timestamp}.{extension}"

def export_companies_to_excel(queryset):
    """Şirketleri Excel'e export et"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Şirketler"
    
    # Header
    headers = ['Şirket Ünvanı', 'Vergi No', 'E-posta', 'IBAN', 'Marka Sayısı', 'Durum', 'Oluşturma']
    ws.append(headers)
    
    # Style header
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    
    # Data
    for company in queryset:
        ws.append([
            company.title,
            company.tax_number,
            company.email,
            company.iban or '-',
            company.brand_count if hasattr(company, 'brand_count') else company.brands.count(),
            'Aktif' if company.is_active else 'Pasif',
            company.created_at.strftime('%Y-%m-%d'),
        ])
    
    # Auto-width
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save
    filename = generate_unique_filename('sirketler', 'xlsx')
    filepath = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    wb.save(filepath)
    
    return f"{settings.MEDIA_URL}exports/{filename}"


def export_companies_to_pdf(queryset):
    """Şirketleri PDF'e export et"""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    
    filename = generate_unique_filename('sirketler', 'pdf')
    filepath = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc = SimpleDocTemplate(filepath, pagesize=landscape(A4))
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#366092'),
        spaceAfter=20
    )
    elements.append(Paragraph("Şirketler Listesi", title_style))
    elements.append(Spacer(1, 12))
    
    # Table
    data = [['Şirket', 'Vergi No', 'E-posta', 'Marka Sayısı', 'Durum']]
    for company in queryset[:100]:
        data.append([
            company.title[:30],
            company.tax_number,
            company.email[:25],
            str(company.brand_count if hasattr(company, 'brand_count') else company.brands.count()),
            'Aktif' if company.is_active else 'Pasif',
        ])
    
    table = Table(data, colWidths=[6*cm, 3*cm, 5*cm, 2.5*cm, 2*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    
    elements.append(table)
    doc.build(elements)
    
    return f"{settings.MEDIA_URL}exports/{filename}"



def export_reports_to_excel(queryset):
    """Raporları Excel'e export et"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Raporlar"
    
    # Header
    headers = ['Başlık', 'Tür', 'Kapsam', 'Tarih', 'Oluşturan', 'Oluşturma Tarihi']
    ws.append(headers)
    
    # Style header
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    
    # Data
    for report in queryset:
        ws.append([
            report.title,
            report.get_report_type_display(),
            report.get_scope_display(),
            report.report_date.strftime('%Y-%m-%d'),
            report.created_by.get_full_name() if report.created_by else '',
            report.created_at.strftime('%Y-%m-%d %H:%M:%S'),
        ])
    
    # Auto-width columns
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save
    filename = generate_unique_filename('raporlar', 'xlsx')
    filepath = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    wb.save(filepath)
    
    return f"{settings.MEDIA_URL}exports/{filename}"


def export_reports_to_pdf(queryset):
    """Raporları PDF'e export et"""
    filename = generate_unique_filename('raporlar', 'pdf')
    filepath = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc = SimpleDocTemplate(filepath, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#366092'),
        alignment=TA_CENTER,
        spaceAfter=30
    )
    elements.append(Paragraph("Raporlar", title_style))
    elements.append(Spacer(1, 12))
    
    # Table data
    data = [['Başlık', 'Tür', 'Kapsam', 'Tarih']]
    for report in queryset[:50]:  # İlk 50 rapor
        data.append([
            report.title[:40],
            report.get_report_type_display(),
            report.get_scope_display(),
            report.report_date.strftime('%Y-%m-%d'),
        ])
    
    # Create table
    table = Table(data, colWidths=[3*inch, 1.5*inch, 1.5*inch, 1.5*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    
    elements.append(table)
    doc.build(elements)
    
    return f"{settings.MEDIA_URL}exports/{filename}"


def export_financial_records_to_excel(queryset):
    """Mali kayıtları Excel'e export et"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Mali Kayıtlar"
    
    # Header
    headers = ['Başlık', 'Tür', 'Tutar', 'Para Birimi', 'Tarih', 'Açıklama', 'Oluşturan']
    ws.append(headers)
    
    # Style header
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2E7D32", end_color="2E7D32", fill_type="solid")
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    
    # Data
    for record in queryset:
        ws.append([
            record.title,
            record.get_type_display(),
            float(record.amount),
            record.currency,
            record.date.strftime('%Y-%m-%d'),
            record.description or '',
            record.created_by.get_full_name() if record.created_by else '',
        ])
    
    # Number format for amount column
    for row in range(2, ws.max_row + 1):
        ws[f'C{row}'].number_format = '#,##0.00'
    
    # Auto-width columns
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save
    filename = generate_unique_filename('mali_kayitlar', 'xlsx')
    filepath = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    wb.save(filepath)
    
    return f"{settings.MEDIA_URL}exports/{filename}"


def export_financial_records_to_pdf(queryset):
    """Mali kayıtları PDF'e export et"""
    filename = generate_unique_filename('mali_kayitlar', 'pdf')
    filepath = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc = SimpleDocTemplate(filepath, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2E7D32'),
        alignment=TA_CENTER,
        spaceAfter=30
    )
    elements.append(Paragraph("Mali Kayıtlar", title_style))
    elements.append(Spacer(1, 12))
    
    # Summary
    total_income = queryset.filter(type='income').aggregate(Sum('amount'))['amount__sum'] or 0
    total_expense = queryset.filter(type='expense').aggregate(Sum('amount'))['amount__sum'] or 0
    
    summary_style = styles['Normal']
    elements.append(Paragraph(f"<b>Toplam Gelir:</b> {total_income:,.2f} TL", summary_style))
    elements.append(Paragraph(f"<b>Toplam Gider:</b> {total_expense:,.2f} TL", summary_style))
    elements.append(Paragraph(f"<b>Net:</b> {(total_income - total_expense):,.2f} TL", summary_style))
    elements.append(Spacer(1, 20))
    
    # Table data
    data = [['Başlık', 'Tür', 'Tutar', 'Tarih']]
    for record in queryset[:50]:  # İlk 50 kayıt
        data.append([
            record.title[:30],
            record.get_type_display(),
            f"{record.amount:,.2f} {record.currency}",
            record.date.strftime('%Y-%m-%d'),
        ])
    
    # Create table
    table = Table(data, colWidths=[3*inch, 1.5*inch, 1.5*inch, 1.5*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E7D32')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('ALIGN', (2, 1), (2, -1), 'RIGHT'),  # Amount column right-aligned
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    
    elements.append(table)
    doc.build(elements)
    
    return f"{settings.MEDIA_URL}exports/{filename}"


def import_financial_records_from_excel(file_path, user):
    """Excel'den mali kayıtları import et"""
    from .models import FinancialRecord, Company, Brand, Branch
    
    try:
        df = pd.read_excel(file_path)
        
        # Required columns
        required_columns = ['title', 'type', 'amount', 'currency', 'date']
        if not all(col in df.columns for col in required_columns):
            return {'success': False, 'error': 'Gerekli sütunlar eksik'}
        
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                # Parse data
                record_data = {
                    'title': row['title'],
                    'type': row['type'],
                    'amount': Decimal(str(row['amount'])),
                    'currency': row['currency'],
                    'date': pd.to_datetime(row['date']).date(),
                    'description': row.get('description', ''),
                    'created_by': user
                }
                
                # İlişkili entity'leri bul (opsiyonel)
                if 'company_tax_number' in row and pd.notna(row['company_tax_number']):
                    try:
                        company = Company.objects.get(tax_number=row['company_tax_number'])
                        record_data['related_company'] = company
                    except Company.DoesNotExist:
                        pass
                
                # Kayıt oluştur
                FinancialRecord.objects.create(**record_data)
                imported_count += 1
                
            except Exception as e:
                errors.append(f"Satır {index + 2}: {str(e)}")
        
        return {
            'success': True,
            'imported_count': imported_count,
            'total_rows': len(df),
            'errors': errors
        }
        
    except Exception as e:
        return {'success': False, 'error': str(e)}


def generate_contract_from_template(template_name, context_data):
    """Şablondan sözleşme oluştur"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    
    # Template path
    template_path = os.path.join(settings.BASE_DIR, 'templates', 'contracts', f'{template_name}.docx')
    
    if not os.path.exists(template_path):
        # Basit bir şablon oluştur
        doc = Document()
        doc.add_heading('Sözleşme', 0)
        doc.add_paragraph(f"Sözleşme No: {context_data.get('contract_number', 'N/A')}")
        doc.add_paragraph(f"Taraflar: {context_data.get('parties', 'N/A')}")
        doc.add_paragraph(f"Başlangıç Tarihi: {context_data.get('start_date', 'N/A')}")
        doc.add_paragraph(f"Bitiş Tarihi: {context_data.get('end_date', 'N/A')}")
    else:
        doc = Document(template_path)
        
        # Replace placeholders
        for paragraph in doc.paragraphs:
            for key, value in context_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in context_data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def mask_sensitive_field(value, field_type='national_id'):
    """Hassas alanları maskele"""
    if not value:
        return None
    
    if field_type == 'national_id' and len(value) == 11:
        return f"{value[:3]}****{value[-2:]}"
    elif field_type == 'iban' and len(value) >= 10:
        return f"{value[:6]}****{value[-4:]}"
    elif field_type == 'phone' and len(value) >= 10:
        return f"{value[:6]}***{value[-2:]}"
    elif field_type == 'email':
        parts = value.split('@')
        if len(parts) == 2:
            return f"{parts[0][:2]}***@{parts[1]}"
    
    return value


def validate_turkish_national_id(national_id):
    """TC Kimlik numarası doğrulama"""
    if not national_id or len(national_id) != 11:
        return False
    
    if not national_id.isdigit():
        return False
    
    if national_id[0] == '0':
        return False
    
    # Checksum algoritması
    digits = [int(d) for d in national_id]
    
    sum_odd = sum(digits[0:9:2])
    sum_even = sum(digits[1:8:2])
    
    check_10 = (sum_odd * 7 - sum_even) % 10
    if check_10 != digits[9]:
        return False
    
    check_11 = sum(digits[:10]) % 10
    if check_11 != digits[10]:
        return False
    
    return True


def calculate_date_difference(start_date, end_date):
    """İki tarih arasındaki farkı hesapla"""
    from datetime import datetime
    
    if isinstance(start_date, str):
        start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
    if isinstance(end_date, str):
        end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
    
    delta = end_date - start_date
    
    years = delta.days // 365
    months = (delta.days % 365) // 30
    days = (delta.days % 365) % 30
    
    return {
        'total_days': delta.days,
        'years': years,
        'months': months,
        'days': days
    }


# Import Sum at the top
from django.db.models import Sum

